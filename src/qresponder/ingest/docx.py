"""DOCX loader (§6, Stage A).

Emits paragraphs and table cells in document order with indices, so questions
buried in prose or tables keep a stable write-back anchor. Heading styles are
flagged as sections.
"""

from __future__ import annotations

from pathlib import Path

import docx as _docx

from .ir import Document, Element


def _is_heading(para) -> bool:
    style = getattr(para, "style", None)
    name = getattr(style, "name", "") or ""
    return name.lower().startswith("heading") or name.lower() == "title"


def _is_bold(para) -> bool:
    runs = para.runs
    return bool(runs) and all((r.bold for r in runs if r.text.strip()))


def load_docx(path: str | Path) -> Document:
    p = Path(path)
    document = _docx.Document(str(p))
    elements: list[Element] = []

    for i, para in enumerate(document.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style: dict = {}
        if _is_heading(para):
            style["section"] = para.style.name
        if _is_bold(para):
            style["bold"] = True
        elements.append(
            Element(
                kind="paragraph",
                location=f"para[{i}]",
                text=text,
                style=style,
            )
        )

    for ti, table in enumerate(document.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue
                elements.append(
                    Element(
                        kind="table_cell",
                        location=f"table[{ti}].r{ri}.c{ci}",
                        text=text,
                        style={"table": True},
                    )
                )

    return Document(source_file=str(p), file_type="docx", elements=elements)
