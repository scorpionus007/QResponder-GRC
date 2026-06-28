"""Generate binary test fixtures (xlsx/docx) deterministically.

Run: python tests/make_fixtures.py
Re-runnable; produces tests/fixtures/sample.xlsx and tests/fixtures/sample.docx.
Kept as a script (not committed binaries are regenerated) so the fixtures carry
the layout signals we test for: merged section headers, bold, fill color.
"""

from pathlib import Path

import openpyxl
from openpyxl.styles import Font, PatternFill
import docx

FIX = Path(__file__).parent / "fixtures"
FIX.mkdir(parents=True, exist_ok=True)


def make_xlsx() -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Security"

    # A merged, colored, bold section banner (color/merge carry meaning).
    ws["A1"] = "Vendor Security Questionnaire — Section 1: Data Protection"
    ws.merge_cells("A1:C1")
    ws["A1"].font = Font(bold=True)
    ws["A1"].fill = PatternFill(start_color="FFD9E1F2", end_color="FFD9E1F2", fill_type="solid")

    # Header row (bold).
    ws["A2"] = "#"
    ws["B2"] = "Question"
    ws["C2"] = "Response"
    for c in ("A2", "B2", "C2"):
        ws[c].font = Font(bold=True)

    rows = [
        ("1", "Do you encrypt data at rest?"),
        ("2", "Do you enforce multi-factor authentication (MFA) for employees?"),
        ("3", "Do you have a documented incident response plan?"),
        ("4", "Please attach your most recent SOC 2 Type II report."),
        ("5", "What is the data retention period for customer backups?"),
    ]
    for i, (num, q) in enumerate(rows, start=3):
        ws[f"A{i}"] = num
        ws[f"B{i}"] = q
        # C column intentionally left blank — that's where answers go.

    wb.save(FIX / "sample.xlsx")
    print(f"wrote {FIX / 'sample.xlsx'}")


def make_docx() -> None:
    d = docx.Document()
    d.add_heading("Vendor Security Questionnaire", level=1)
    d.add_heading("Section 1: Access Control", level=2)
    d.add_paragraph("Please answer the following questions about your controls.")
    d.add_paragraph("Do you enforce multi-factor authentication (MFA) for employees?")
    d.add_paragraph("Do you have a documented incident response plan?")
    d.save(FIX / "sample.docx")
    print(f"wrote {FIX / 'sample.docx'}")


if __name__ == "__main__":
    make_xlsx()
    make_docx()
