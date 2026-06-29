"""Format-perfect write-back tests (C3, §15). Handles the openpyxl traps."""

from pathlib import Path

import openpyxl

from qresponder.models import (
    AnswerResult,
    AnswerType,
    Confidence,
    QuestionnaireResult,
    ReviewReason,
    Status,
)
from qresponder.output.writeback import write_back


def _answered(qid, text, answer, anchor):
    return AnswerResult(
        question_id=qid,
        question_text=text,
        answer=answer,
        answer_type=AnswerType.YES_NO,
        confidence=Confidence.HIGH,
        status=Status.ANSWERED,
        answer_location_hint=anchor,
    )


def test_writeback_writes_to_merged_anchor_without_raising(tmp_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q"
    ws["A1"] = "Question"
    ws["B1"] = "Response"
    ws["A2"] = "Do you encrypt data at rest?"
    ws.merge_cells("B2:C2")  # the answer cell is a merged range (anchor B2)
    orig = tmp_path / "q.xlsx"
    wb.save(orig)

    # Anchor points at C2 — a NON-anchor member of the merge. Write-back must
    # redirect to the top-left B2 and not raise.
    result = QuestionnaireResult(
        source_file=str(orig),
        results=[_answered("q1", "Do you encrypt data at rest?", "Yes, AES-256.", "Q!C2")],
    )
    info = write_back(result, str(orig), str(tmp_path / "out"))

    assert info["written"], info
    out_wb = openpyxl.load_workbook(info["written"])
    assert out_wb["Q"]["B2"].value == "Yes, AES-256."  # landed on the anchor

    # The ORIGINAL is never modified.
    assert openpyxl.load_workbook(orig)["Q"]["B2"].value is None
    assert Path(info["written"]).name == "q.answered.xlsx"


def test_writeback_preserves_style(tmp_path):
    from openpyxl.styles import Font

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q"
    ws["A1"] = "Question"
    ws["A2"] = "Encrypt at rest?"
    ws["B2"].font = Font(bold=True)  # pre-existing style on the answer cell
    orig = tmp_path / "q.xlsx"
    wb.save(orig)

    result = QuestionnaireResult(
        source_file=str(orig),
        results=[_answered("q1", "Encrypt at rest?", "Yes.", "Q!B2")],
    )
    info = write_back(result, str(orig), str(tmp_path / "out"))
    out_ws = openpyxl.load_workbook(info["written"])["Q"]
    assert out_ws["B2"].value == "Yes."
    assert out_ws["B2"].font.bold is True  # style preserved (value-only write)


def test_writeback_falls_back_on_images(tmp_path):
    from openpyxl.drawing.image import Image
    from PIL import Image as PILImage

    img = tmp_path / "px.png"
    PILImage.new("RGB", (2, 2), (255, 0, 0)).save(img)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q"
    ws["A2"] = "Encrypt at rest?"
    ws.add_image(Image(str(img)), "D1")
    orig = tmp_path / "q.xlsx"
    wb.save(orig)

    result = QuestionnaireResult(
        source_file=str(orig),
        results=[_answered("q1", "Encrypt at rest?", "Yes.", "Q!B2")],
    )
    info = write_back(result, str(orig), str(tmp_path / "out"))
    # Don't risk dropping the user's image — fall back, don't write.
    assert info["fallback"] is True
    assert info["written"] is None
    assert not (tmp_path / "out" / "q.answered.xlsx").exists()


def test_writeback_heuristic_response_column(tmp_path):
    """No explicit answer anchor -> find the 'Response' column in the question's row."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q"
    ws["A1"] = "Question"
    ws["B1"] = "Response"
    ws["A2"] = "Encrypt at rest?"
    orig = tmp_path / "q.xlsx"
    wb.save(orig)

    r = _answered("q1", "Encrypt at rest?", "Yes, AES-256.", None)
    r.location_hint = "Q!A2"  # only the question cell is known
    result = QuestionnaireResult(source_file=str(orig), results=[r])
    info = write_back(result, str(orig), str(tmp_path / "out"))
    out_ws = openpyxl.load_workbook(info["written"])["Q"]
    assert out_ws["B2"].value == "Yes, AES-256."


def test_writeback_never_overwrites_filled_response_cell(tmp_path):
    """SH3: if the header-matched response cell already has content, write-back
    must not overwrite it — it falls through to the next empty cell."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q"
    ws["A1"] = "Question"
    ws["B1"] = "Response"
    ws["A2"] = "Encrypt at rest?"
    ws["B2"] = "PRE-EXISTING ANSWER"  # response column already filled in this row
    orig = tmp_path / "q.xlsx"
    wb.save(orig)

    r = _answered("q1", "Encrypt at rest?", "Yes, AES-256.", None)
    r.location_hint = "Q!A2"
    result = QuestionnaireResult(source_file=str(orig), results=[r])
    info = write_back(result, str(orig), str(tmp_path / "out"))

    out_ws = openpyxl.load_workbook(info["written"])["Q"]
    assert out_ws["B2"].value == "PRE-EXISTING ANSWER"  # not overwritten
    # The answer landed in the next empty cell instead.
    assert out_ws["C2"].value == "Yes, AES-256."


def test_writeback_preserves_data_validation_dropdown(tmp_path):
    """Part F: a dropdown / data-validation answer cell keeps its validation
    after write-back, and the written value lands in the cell."""
    from openpyxl.worksheet.datavalidation import DataValidation

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q"
    ws["A1"] = "Question"
    ws["B1"] = "Response"
    ws["A2"] = "Do you encrypt at rest?"
    dv = DataValidation(type="list", formula1='"Yes,No"', allow_blank=True)
    ws.add_data_validation(dv)
    dv.add("B2")  # dropdown on the answer cell
    orig = tmp_path / "q.xlsx"
    wb.save(orig)

    result = QuestionnaireResult(
        source_file=str(orig),
        results=[_answered("q1", "Do you encrypt at rest?", "Yes", "Q!B2")],
    )
    info = write_back(result, str(orig), str(tmp_path / "out"))
    out_wb = openpyxl.load_workbook(info["written"])
    out_ws = out_wb["Q"]
    assert out_ws["B2"].value == "Yes"
    # The data validation survived the round-trip.
    sqrefs = " ".join(str(d.sqref) for d in out_ws.data_validations.dataValidation)
    assert "B2" in sqrefs
    assert any(d.type == "list" for d in out_ws.data_validations.dataValidation)


def test_writeback_coerces_to_dropdown_option(tmp_path):
    """Part F: a verbose yes/no answer is coerced to the dropdown's allowed value."""
    from openpyxl.worksheet.datavalidation import DataValidation

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q"
    ws["A2"] = "Encrypt at rest?"
    dv = DataValidation(type="list", formula1='"Yes,No"', allow_blank=True)
    ws.add_data_validation(dv)
    dv.add("B2")
    orig = tmp_path / "q.xlsx"
    wb.save(orig)

    result = QuestionnaireResult(
        source_file=str(orig),
        results=[_answered("q1", "Encrypt at rest?", "Yes. All data is AES-256 encrypted.", "Q!B2")],
    )
    info = write_back(result, str(orig), str(tmp_path / "out"))
    out_ws = openpyxl.load_workbook(info["written"])["Q"]
    assert out_ws["B2"].value == "Yes"  # coerced to the allowed dropdown option


def _flagged(qid, text, anchor):
    return AnswerResult(
        question_id=qid, question_text=text, answer="", answer_type=AnswerType.YES_NO,
        confidence=Confidence.LOW, status=Status.NEEDS_REVIEW,
        review_reason=ReviewReason.UNSUPPORTED, answer_location_hint=anchor,
        missing_info="Not supported by the KB.",
    )


def test_review_markers_default_and_toggle_and_safety(tmp_path):
    """Part C: NEEDS_REVIEW cells get a visible marker by default; ANSWERED cells
    untouched; --no-review-markers restores blank; never overwrites a filled cell."""
    from qresponder.models import ReviewReason

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Q"
    ws["A2"] = "Encrypt at rest?"          # answered -> filled
    ws["A3"] = "Unsupported?"              # flagged
    ws["A4"] = "Already filled?"           # flagged but answer cell pre-filled
    ws["B4"] = "PRE-EXISTING"
    orig = tmp_path / "q.xlsx"
    wb.save(orig)

    results = [
        _answered("q1", "Encrypt at rest?", "Yes", "Q!B2"),
        _flagged("q2", "Unsupported?", "Q!B3"),
        _flagged("q3", "Already filled?", "Q!B4"),
    ]
    res = QuestionnaireResult(source_file=str(orig), results=results)

    # Default: markers on.
    info = write_back(res, str(orig), str(tmp_path / "out"))
    ws_o = openpyxl.load_workbook(info["written"])["Q"]
    assert ws_o["B2"].value == "Yes"                       # answered untouched
    assert ws_o["B3"].value.startswith("⚠ NEEDS REVIEW")   # flagged marked + reason
    assert "unsupported" in ws_o["B3"].value.lower()
    assert ws_o["B4"].value == "PRE-EXISTING"              # never overwrote a filled cell

    # Toggle off: NEEDS_REVIEW cell stays blank.
    info2 = write_back(res, str(orig), str(tmp_path / "out2"), review_markers=False)
    ws2 = openpyxl.load_workbook(info2["written"])["Q"]
    assert ws2["B3"].value in (None, "")
    assert ws2["B2"].value == "Yes"

    # The original file is never modified.
    orig_ws = openpyxl.load_workbook(orig)["Q"]
    assert orig_ws["B3"].value in (None, "")


def test_answered_xlsx_shows_marker(tmp_path):
    from qresponder.output.writer import write_xlsx

    res = QuestionnaireResult(source_file="q.xlsx", results=[_flagged("q1", "Unsupported?", None)])
    p = write_xlsx(res, tmp_path / "answered.xlsx")
    cell = openpyxl.load_workbook(p).active.cell(row=2, column=3).value
    assert cell.startswith("⚠ NEEDS REVIEW")


def test_writeback_docx_paragraph(tmp_path):
    import docx

    d = docx.Document()
    d.add_paragraph("Do you encrypt at rest?")  # para[0]
    orig = tmp_path / "q.docx"
    d.save(orig)

    r = _answered("q1", "Do you encrypt at rest?", "Yes, AES-256.", "para[0]")
    r.answer_type = AnswerType.TEXT
    result = QuestionnaireResult(source_file=str(orig), results=[r])
    info = write_back(result, str(orig), str(tmp_path / "out"))
    out_doc = docx.Document(info["written"])
    assert "AES-256" in out_doc.paragraphs[0].text
